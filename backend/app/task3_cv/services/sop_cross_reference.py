# SOP Cross-Reference Service for P&ID Analysis (Task 3)
#
# This is DIFFERENT from Task 2's SOP compliance checking:
# - Task 2: Checks if SOP procedures comply with regulations
# - Task 3: Extracts component tags from SOP and cross-references with P&ID diagram
#
# Uses LLM (OpenAI) for intelligent parsing of SOP documents
#
from typing import List, Dict, Optional, Tuple
import networkx as nx
import re
import json
from docx import Document
from pathlib import Path
from openai import OpenAI

from app.config import settings


class SOPComponentExtractor:
    """
    Extracts equipment/component references from SOP documents using LLM.

    Uses OpenAI to intelligently parse SOP documents and extract:
    - Equipment tags (V-720, P-101, E-742, etc.)
    - Component types (valve, pump, exchanger, etc.)
    - Operating limits (pressure, temperature)
    - Context and descriptions
    """

    # Component type synonyms for matching
    COMPONENT_SYNONYMS = {
        "valve": ["valve", "v-", "vv-", "control valve", "shut-off", "gate", "ball", "check", "stabilizer"],
        "pump": ["pump", "p-", "pp-", "centrifugal", "booster"],
        "tank": ["tank", "t-", "tk-", "vessel", "drum", "storage", "tower"],
        "sensor": ["sensor", "transmitter", "indicator", "gauge"],
        "heat_exchanger": ["exchanger", "e-", "he-", "cooler", "heater", "htr"],
        "compressor": ["compressor", "c-"],
        "filter": ["filter", "f-", "strainer", "particulate"],
        "cooler": ["cooler", "ac-", "after cooler"],
    }

    def __init__(self):
        self.client = None
        if settings.OPENAI_API_KEY:
            self.client = OpenAI(api_key=settings.OPENAI_API_KEY)

    async def parse_sop(self, sop_path: str) -> Dict:
        """
        Parse SOP document using LLM and extract component references.

        Returns dict with:
        - filename: SOP filename
        - sections: List of sections with their components
        - all_components: Deduplicated list of all components found
        """
        path = Path(sop_path)

        # Extract raw text and table data from document
        if path.suffix.lower() in {".docx", ".doc"}:
            doc_content = self._extract_docx_content(path)
        elif path.suffix.lower() == ".pdf":
            doc_content = self._extract_pdf_content(path)
        else:
            raise ValueError(f"Unsupported file type: {path.suffix}")

        # Use LLM to parse the content
        if self.client:
            components = await self._llm_extract_components(doc_content)
        else:
            print("[SOP Parser] OpenAI not configured, using empty result")
            components = []

        return {
            "filename": path.name,
            "sections": [{"title": "Full Document", "content": doc_content, "components": components}],
            "all_components": components,
        }

    def _extract_docx_content(self, path: Path) -> str:
        """Extract all text content from DOCX including tables."""
        print(f"[SOP Parser] Reading DOCX: {path}")
        doc = Document(str(path))
        content_parts = []

        # Extract paragraphs
        para_count = 0
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                content_parts.append(text)
                para_count += 1

        print(f"[SOP Parser] Extracted {para_count} paragraphs")

        # Extract tables - this is where Design Limits table should be
        table_count = 0
        for table in doc.tables:
            table_count += 1
            content_parts.append("\n[TABLE]")
            row_count = 0
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                row_text = " | ".join(cells)
                content_parts.append(row_text)
                row_count += 1
            content_parts.append("[/TABLE]\n")
            print(f"[SOP Parser] Table {table_count}: {row_count} rows")

        print(f"[SOP Parser] Extracted {table_count} tables total")

        full_content = "\n".join(content_parts)
        print(f"[SOP Parser] Total content: {len(full_content)} chars")

        # Log first 500 chars for debugging
        print(f"[SOP Parser] Content preview: {full_content[:500]}...")

        return full_content

    def _extract_pdf_content(self, path: Path) -> str:
        """Extract text content from PDF."""
        import fitz  # PyMuPDF

        doc = fitz.open(str(path))
        full_text = ""

        for page in doc:
            full_text += page.get_text("text") + "\n"

        doc.close()
        return full_text

    async def _llm_extract_components(self, content: str) -> List[Dict]:
        """Use LLM to extract equipment components from SOP content."""

        # Log content length for debugging
        print(f"[SOP Parser] Document content length: {len(content)} chars")
        if len(content) < 50:
            print(f"[SOP Parser] WARNING: Content seems too short: {content[:200]}")

        prompt = f"""Extract equipment from "Design Limits" table in SOP. Sub-parts (Shell/Tube) = separate entries.

Extract: tag, description, pressure, temperature

Example:
| F-715 A and B Particulate Filters | 275 | 100 |
| E-742 Exchanger (Shell) | 300 | 375 |
→ {{"components": [{{"tag": "F-715", "description": "Particulate Filters", "pressure": "275", "temperature": "100"}}, {{"tag": "E-742", "description": "Exchanger (Shell)", "pressure": "300", "temperature": "375"}}]}}

SOP:
{content}

Return: {{"components": [...]}}"""

        try:
            response = self.client.chat.completions.create(
                model="gpt-5-nano",
                messages=[
                    {"role": "system", "content": "SOP parser. Return only JSON."},
                    {"role": "user", "content": prompt}
                ],
                response_format={"type": "json_object"}
            )

            result_text = response.choices[0].message.content
            print(f"[SOP Parser] LLM response: {result_text[:500]}...")

            # Parse JSON response
            result = json.loads(result_text)

            # Handle various response formats
            if isinstance(result, list):
                components = result
            elif isinstance(result, dict):
                # Try common keys
                components = result.get("components") or result.get("equipment") or result.get("items") or []
            else:
                components = []

            print(f"[SOP Parser] LLM extracted {len(components)} components")

            # Log each component for debugging
            for comp in components[:5]:
                print(f"  - {comp.get('tag')}: {comp.get('pressure')} psig, {comp.get('temperature')}°F")

            return components

        except Exception as e:
            print(f"[SOP Parser] LLM extraction error: {e}")
            import traceback
            traceback.print_exc()
            return []


class SOPCrossReference:
    """
    Cross-references P&ID extracted specifications with SOP design limits.

    The key comparison:
    - P&ID title blocks contain: DESIGN: 275 PSIG @ 100°F
    - SOP contains design limits table: F-715: 275 psig, 100°F

    Identifies:
    - Matches: P&ID spec matches SOP limit
    - Pressure discrepancies: P&ID design pressure != SOP limit
    - Temperature discrepancies: P&ID design temp != SOP limit
    - Missing: Components in one but not the other
    """

    def __init__(self):
        self.extractor = SOPComponentExtractor()

    async def parse_sop(self, sop_path: str) -> Dict:
        """Parse SOP and extract component references using LLM."""
        return await self.extractor.parse_sop(sop_path)

    async def cross_reference_with_specs(
        self,
        pid_specs: List[Dict],
        sop_data: Dict
    ) -> Dict:
        """
        Cross-reference P&ID extracted specs with SOP design limits using LLM.

        Args:
            pid_specs: Equipment specs extracted from P&ID title blocks
            sop_data: Parsed SOP data with component limits

        Returns:
            Detailed comparison report with matches and discrepancies
        """
        sop_components = sop_data.get("all_components", [])

        print(f"[CrossRef] P&ID specs: {[s.get('tag') for s in pid_specs]}")
        print(f"[CrossRef] SOP components: {[c.get('tag') for c in sop_components]}")

        # Use LLM to do intelligent matching and comparison
        if self.extractor.client and pid_specs and sop_components:
            results = await self._llm_cross_reference(pid_specs, sop_components)
        else:
            print("[CrossRef] No LLM client or empty data, using basic matching")
            results = self._basic_cross_reference(pid_specs, sop_components)

        # Generate summary
        results["summary"] = self._generate_comparison_summary(results)

        return results

    async def _llm_cross_reference(
        self,
        pid_specs: List[Dict],
        sop_components: List[Dict]
    ) -> Dict:
        """Use LLM to intelligently match and compare P&ID specs with SOP limits."""

        # Format data for LLM (compact JSON)
        pid_text = json.dumps(pid_specs)
        sop_text = json.dumps(sop_components)

        prompt = f"""Cross-reference P&ID specs with SOP limits. Match by tag (F-715 = F-715-A/B, E-742-SHELL = E-742 (Shell)). Tolerance: ±5 psig/°F.

P&ID: {pid_text}
SOP: {sop_text}

Return JSON:
{{"matches": [{{"tag":"F-715","sop_description":"...","pid_description":"...","sop_pressure":275,"pid_pressure":275,"sop_temperature":100,"pid_temperature":100,"status":"match","pressure_issue":null,"temperature_issue":null}}],
"pressure_discrepancies": [...],
"temperature_discrepancies": [...],
"missing_in_pid": [{{"tag":"X","sop_description":"...","sop_pressure":"...","sop_temperature":"...","issue":"Not found in P&ID"}}],
"missing_in_sop": [...]}}"""

        try:
            response = self.extractor.client.chat.completions.create(
                model="gpt-5-nano",
                messages=[
                    {"role": "system", "content": "Equipment comparison. Return only JSON."},
                    {"role": "user", "content": prompt}
                ],
                response_format={"type": "json_object"}
            )

            result_text = response.choices[0].message.content
            print(f"[CrossRef] LLM response: {result_text[:500]}...")

            result = json.loads(result_text)

            # Ensure all required keys exist
            return {
                "comparisons": result.get("comparisons", []),
                "matches": result.get("matches", []),
                "pressure_discrepancies": result.get("pressure_discrepancies", []),
                "temperature_discrepancies": result.get("temperature_discrepancies", []),
                "missing_in_pid": result.get("missing_in_pid", []),
                "missing_in_sop": result.get("missing_in_sop", []),
            }

        except Exception as e:
            print(f"[CrossRef] LLM error: {e}")
            import traceback
            traceback.print_exc()
            return self._basic_cross_reference(pid_specs, sop_components)

    def _basic_cross_reference(
        self,
        pid_specs: List[Dict],
        sop_components: List[Dict]
    ) -> Dict:
        """Fallback basic matching without LLM."""
        results = {
            "comparisons": [],
            "matches": [],
            "pressure_discrepancies": [],
            "temperature_discrepancies": [],
            "missing_in_pid": [],
            "missing_in_sop": [],
        }

        # Build lookup
        pid_by_tag = {}
        for spec in pid_specs:
            tag = spec.get("tag", "").upper()
            part = spec.get("part", "")
            key = f"{tag}-{part.upper()}" if part else tag
            pid_by_tag[key] = spec

        matched_pid_tags = set()

        for sop_comp in sop_components:
            sop_tag = sop_comp.get("tag", "").upper()
            desc = sop_comp.get("description", "")

            # Try to find match
            pid_spec = None
            for pid_key, spec in pid_by_tag.items():
                if sop_tag in pid_key or pid_key in sop_tag:
                    pid_spec = spec
                    matched_pid_tags.add(pid_key)
                    break

            if pid_spec:
                comparison = self._compare_specs(sop_tag, sop_comp, pid_spec)
                results["comparisons"].append(comparison)
                if comparison["status"] == "match":
                    results["matches"].append(comparison)
                else:
                    if comparison.get("pressure_issue"):
                        results["pressure_discrepancies"].append(comparison)
                    if comparison.get("temperature_issue"):
                        results["temperature_discrepancies"].append(comparison)
            else:
                results["missing_in_pid"].append({
                    "tag": sop_tag,
                    "sop_description": desc,
                    "issue": "Not found in P&ID"
                })

        # Find unmatched P&ID specs
        for pid_key, spec in pid_by_tag.items():
            if pid_key not in matched_pid_tags:
                results["missing_in_sop"].append({
                    "tag": pid_key,
                    "pid_description": spec.get("description", ""),
                    "issue": "Not in SOP"
                })

        return results

    def _find_matching_spec(self, tag: str, specs_dict: Dict) -> Optional[Dict]:
        """Find matching spec by tag, supporting partial matches and sub-parts."""
        # Exact match first
        if tag in specs_dict:
            return specs_dict[tag]

        # Try matching with common variations
        tag_upper = tag.upper()

        for spec_tag, spec in specs_dict.items():
            spec_tag_upper = spec_tag.upper()

            # Exact match after normalization
            if tag_upper == spec_tag_upper:
                return spec

            # Extract base tag (without part suffix like -SHELL, -TUBE)
            base_tag = tag_upper.split('-SHELL')[0].split('-TUBE')[0].rstrip('-')
            base_spec_tag = spec_tag_upper.split('-SHELL')[0].split('-TUBE')[0].rstrip('-')

            # Check if base tags match (handles E-742-SHELL matching E-742)
            if base_tag == base_spec_tag:
                return spec

            # Also strip A/B suffixes for matching
            base_tag = base_tag.rstrip('AB').rstrip('-')
            base_spec_tag = base_spec_tag.rstrip('AB').rstrip('-')
            if base_tag == base_spec_tag:
                return spec
            if tag in spec_tag or spec_tag in tag:
                return spec

        return None

    def _compare_specs(self, tag: str, sop_comp: Dict, pid_spec: Dict) -> Dict:
        """Compare SOP limits against P&ID specs."""
        comparison = {
            "tag": tag,
            "sop_description": sop_comp.get("description", ""),
            "pid_description": pid_spec.get("description", ""),
            "status": "match",  # Will be updated if discrepancies found
            "pressure_issue": None,
            "temperature_issue": None,
        }

        # Parse SOP pressure (e.g., "275 psig" -> 275)
        sop_pressure = self._parse_pressure(sop_comp.get("pressure", ""))
        pid_pressure = pid_spec.get("design_pressure")

        comparison["sop_pressure"] = sop_pressure
        comparison["pid_pressure"] = pid_pressure

        # Compare pressures
        if sop_pressure is not None and pid_pressure is not None:
            if abs(sop_pressure - pid_pressure) > 1:  # Allow 1 psi tolerance
                comparison["status"] = "discrepancy"
                comparison["pressure_issue"] = f"P&ID shows {pid_pressure} psig, SOP requires {sop_pressure} psig"

        # Parse SOP temperature (e.g., "100°F" -> 100)
        sop_temp = self._parse_temperature(sop_comp.get("temperature", ""))
        pid_temp = pid_spec.get("design_temperature")

        comparison["sop_temperature"] = sop_temp
        comparison["pid_temperature"] = pid_temp

        # Compare temperatures
        if sop_temp is not None and pid_temp is not None:
            if abs(sop_temp - pid_temp) > 1:  # Allow 1°F tolerance
                comparison["status"] = "discrepancy"
                comparison["temperature_issue"] = f"P&ID shows {pid_temp}°F, SOP requires {sop_temp}°F"

        return comparison

    def _parse_pressure(self, pressure_str: str) -> Optional[float]:
        """Parse pressure string to float value."""
        if not pressure_str:
            return None
        import re
        match = re.search(r'(-?\d+\.?\d*)', str(pressure_str))
        if match:
            return float(match.group(1))
        return None

    def _parse_temperature(self, temp_str: str) -> Optional[float]:
        """Parse temperature string to float value."""
        if not temp_str:
            return None
        import re
        # Handle ranges like "-20 to 400"
        if "to" in str(temp_str).lower():
            # Use the upper limit for comparison
            match = re.search(r'to\s*(-?\d+\.?\d*)', str(temp_str))
            if match:
                return float(match.group(1))
        match = re.search(r'(-?\d+\.?\d*)', str(temp_str))
        if match:
            return float(match.group(1))
        return None

    def _generate_comparison_summary(self, results: Dict) -> Dict:
        """Generate summary of comparison results."""
        total_compared = len(results["comparisons"])
        matches = len(results["matches"])
        pressure_issues = len(results["pressure_discrepancies"])
        temp_issues = len(results["temperature_discrepancies"])

        return {
            "total_components_compared": total_compared,
            "matches": matches,
            "pressure_discrepancies": pressure_issues,
            "temperature_discrepancies": temp_issues,
            "missing_in_pid": len(results["missing_in_pid"]),
            "missing_in_sop": len(results["missing_in_sop"]),
            "match_rate": round(matches / max(total_compared, 1) * 100, 1),
            "compliance_status": "PASS" if (pressure_issues == 0 and temp_issues == 0) else "REVIEW REQUIRED",
        }

    # Keep legacy method for backwards compatibility with graph-based comparison
    async def cross_reference(
        self,
        graph: nx.DiGraph,
        sop_data: Dict
    ) -> Dict:
        """Legacy cross-reference using graph nodes (backwards compatible)."""
        results = {
            "matches": [],
            "missing_in_pid": [],
            "missing_in_sop": [],
            "type_mismatches": [],
            "connection_issues": [],
        }

        # Build lookup dicts
        pid_components = {
            node: data for node, data in graph.nodes(data=True)
        }

        # Build SOP components lookup
        sop_components = {}
        for comp in sop_data.get("all_components", []):
            tag = comp.get("tag", "")
            if tag:
                sop_components[tag] = comp

        print(f"[CrossRef] P&ID components: {list(pid_components.keys())}")
        print(f"[CrossRef] SOP components: {list(sop_components.keys())}")

        # Find matches and mismatches
        for tag, sop_comp in sop_components.items():
            matched_pid_tag = None
            if tag in pid_components:
                matched_pid_tag = tag
            else:
                for pid_tag in pid_components:
                    if tag in pid_tag or pid_tag in tag:
                        matched_pid_tag = pid_tag
                        break

            if matched_pid_tag:
                pid_comp = pid_components[matched_pid_tag]
                sop_type = sop_comp.get("type", "")
                pid_type = pid_comp.get("type", "")

                if self._types_match(sop_type, pid_type):
                    results["matches"].append({
                        "tag": tag,
                        "sop_type": sop_type,
                        "pid_type": pid_type,
                        "description": sop_comp.get("description", ""),
                        "pressure": sop_comp.get("pressure", ""),
                        "temperature": sop_comp.get("temperature", ""),
                        "status": "match",
                    })
                else:
                    results["type_mismatches"].append({
                        "tag": tag,
                        "sop_type": sop_type,
                        "pid_type": pid_type,
                        "description": sop_comp.get("description", ""),
                        "issue": f"Type mismatch: SOP says '{sop_type}', P&ID shows '{pid_type}'",
                    })
            else:
                results["missing_in_pid"].append({
                    "tag": tag,
                    "type": sop_comp.get("type", ""),
                    "description": sop_comp.get("description", ""),
                    "pressure": sop_comp.get("pressure", ""),
                    "temperature": sop_comp.get("temperature", ""),
                    "context": sop_comp.get("context", ""),
                    "issue": "Component mentioned in SOP but not found in P&ID",
                })

        for tag, pid_comp in pid_components.items():
            found_in_sop = False
            for sop_tag in sop_components:
                if tag in sop_tag or sop_tag in tag:
                    found_in_sop = True
                    break

            if not found_in_sop:
                results["missing_in_sop"].append({
                    "tag": tag,
                    "type": pid_comp.get("type", ""),
                    "confidence": pid_comp.get("confidence", 0),
                    "issue": "Component detected in P&ID but not documented in SOP",
                })

        results["summary"] = self._generate_summary(results)
        return results

    def _types_match(self, sop_type: str, pid_type: str) -> bool:
        """Check if component types match (considering synonyms)."""
        if not sop_type or not pid_type:
            return True  # Can't compare if one is missing

        sop_type = sop_type.lower().strip()
        pid_type = pid_type.lower().strip()

        if sop_type == pid_type:
            return True

        # Direct substring match
        if sop_type in pid_type or pid_type in sop_type:
            return True

        # Check synonyms
        for base_type, synonyms in SOPComponentExtractor.COMPONENT_SYNONYMS.items():
            sop_matches = sop_type == base_type or any(s in sop_type for s in synonyms)
            pid_matches = pid_type == base_type or any(s in pid_type for s in synonyms)
            if sop_matches and pid_matches:
                return True

        return False

    def _generate_summary(self, results: Dict) -> Dict:
        """Generate summary of cross-reference results."""
        total_sop = len(results["matches"]) + len(results["missing_in_pid"])
        total_pid = len(results["matches"]) + len(results["missing_in_sop"])

        return {
            "total_sop_components": total_sop,
            "total_pid_components": total_pid,
            "matched_count": len(results["matches"]),
            "missing_in_pid_count": len(results["missing_in_pid"]),
            "missing_in_sop_count": len(results["missing_in_sop"]),
            "type_mismatch_count": len(results["type_mismatches"]),
            "connection_issue_count": len(results.get("connection_issues", [])),
            "match_rate": round(
                len(results["matches"]) / max(total_sop, 1) * 100, 1
            ),
        }


# Singleton instances
sop_extractor = SOPComponentExtractor()
sop_cross_reference = SOPCrossReference()
