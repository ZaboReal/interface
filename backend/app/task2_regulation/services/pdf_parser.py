# Document Parser Service for PDF and DOCX files
# Uses Unstructured.io when available, falls back to PyMuPDF
import fitz  # PyMuPDF
from docx import Document
from pathlib import Path
from typing import List, Dict
import re
import asyncio

from app.config import settings


class DocumentParser:
    """Parses PDF and DOCX documents and extracts structured text."""

    def __init__(self):
        self.supported_extensions = {".pdf", ".docx", ".doc"}
        self._unstructured_client = None
        self._unstructured_initialized = False
        self.pages_per_section = 1  # Process 1 page at a time for better extraction

    def _init_unstructured(self):
        """Initialize Unstructured.io client if configured."""
        if self._unstructured_initialized:
            return self._unstructured_client is not None

        if settings.UNSTRUCTURED_API_KEY and settings.UNSTRUCTURED_API_KEY != "your-unstructured-api-key-here":
            try:
                import unstructured_client
                self._unstructured_client = unstructured_client.UnstructuredClient(
                    api_key_auth=settings.UNSTRUCTURED_API_KEY
                )
                print("[DocumentParser] Unstructured.io API configured")
            except Exception as e:
                print(f"[DocumentParser] Unstructured.io not available: {e}")
                self._unstructured_client = None

        self._unstructured_initialized = True
        return self._unstructured_client is not None

    def parse(self, file_path: str) -> Dict:
        """Parse a document and return structured content."""
        path = Path(file_path)
        extension = path.suffix.lower()

        if extension == ".pdf":
            return self._parse_pdf(path)
        elif extension in {".docx", ".doc"}:
            return self._parse_docx(path)
        else:
            raise ValueError(f"Unsupported file type: {extension}")

    async def parse_async(self, file_path: str) -> Dict:
        """Async parse using Unstructured.io when available."""
        path = Path(file_path)
        extension = path.suffix.lower()

        if extension == ".pdf":
            if self._init_unstructured():
                try:
                    return await self._parse_pdf_unstructured(path)
                except Exception as e:
                    print(f"[DocumentParser] Unstructured failed, using PyMuPDF: {e}")
                    return self._parse_pdf(path)
            return self._parse_pdf(path)
        elif extension in {".docx", ".doc"}:
            return self._parse_docx(path)
        else:
            raise ValueError(f"Unsupported file type: {extension}")

    async def _parse_pdf_unstructured(self, path: Path) -> Dict:
        """Parse PDF using Unstructured.io API."""
        from unstructured_client.models import shared

        print(f"[DocumentParser] Parsing with Unstructured.io: {path.name}")

        with open(path, "rb") as f:
            file_content = f.read()

        req = {
            "partition_parameters": {
                "files": {
                    "content": file_content,
                    "file_name": path.name,
                },
                "strategy": shared.Strategy.AUTO,
                "languages": ["eng"],
                "split_pdf_page": True,
                "split_pdf_allow_failed": True,
                "split_pdf_concurrency_level": 15,
            }
        }

        res = await self._unstructured_client.general.partition_async(request=req)
        elements = res.elements if hasattr(res, 'elements') else []

        # Process elements by page
        pages_dict: Dict[int, List[str]] = {}
        full_text_parts = []

        for element in elements:
            if hasattr(element, 'to_dict'):
                elem_dict = element.to_dict()
            elif isinstance(element, dict):
                elem_dict = element
            else:
                continue

            text = elem_dict.get("text", "")
            metadata = elem_dict.get("metadata", {})
            page_number = metadata.get("page_number", 1)

            if text.strip():
                if page_number not in pages_dict:
                    pages_dict[page_number] = []
                pages_dict[page_number].append(text)
                full_text_parts.append(text)

        # Build pages list
        pages = []
        for page_num in sorted(pages_dict.keys()):
            pages.append({
                "page_number": page_num,
                "text": "\n".join(pages_dict[page_num]),
            })

        page_count = max(pages_dict.keys()) if pages_dict else 0
        full_text = "\n\n".join(full_text_parts)

        # Create 10-page sections
        sections = self._create_page_sections(pages)

        return {
            "filename": path.name,
            "type": "pdf",
            "page_count": page_count,
            "pages": pages,
            "full_text": full_text,
            "sections": sections,
            "parser": "unstructured",
        }

    def _parse_pdf(self, path: Path) -> Dict:
        """Extract text and structure from PDF using PyMuPDF."""
        print(f"[DocumentParser] Parsing with PyMuPDF: {path.name}")

        doc = fitz.open(str(path))
        pages = []
        full_text = []

        for page_num, page in enumerate(doc):
            text = page.get_text("text")
            pages.append({
                "page_number": page_num + 1,
                "text": text,
            })
            full_text.append(text)

        doc.close()

        combined_text = "\n\n".join(full_text)

        # Create 10-page sections
        sections = self._create_page_sections(pages)

        return {
            "filename": path.name,
            "type": "pdf",
            "page_count": len(pages),
            "pages": pages,
            "full_text": combined_text,
            "sections": sections,
            "parser": "pymupdf",
        }

    def _create_page_sections(self, pages: List[Dict]) -> List[Dict]:
        """
        Group pages into sections of N pages each for chunked processing.
        Each section contains up to `pages_per_section` pages.
        """
        sections = []

        for i in range(0, len(pages), self.pages_per_section):
            chunk = pages[i:i + self.pages_per_section]
            if not chunk:
                continue

            start_page = chunk[0].get("page_number", i + 1)
            end_page = chunk[-1].get("page_number", i + len(chunk))

            # Combine text from all pages in this section
            section_text = "\n\n".join([p.get("text", "") for p in chunk])

            sections.append({
                "title": f"Pages {start_page}-{end_page}",
                "start_page": start_page,
                "end_page": end_page,
                "content": section_text,
                "page_count": len(chunk),
            })

        print(f"[DocumentParser] Created {len(sections)} sections ({self.pages_per_section} pages each)")
        return sections

    def _parse_docx(self, path: Path) -> Dict:
        """Extract text and structure from DOCX."""
        print(f"[DocumentParser] Parsing DOCX: {path.name}")

        doc = Document(str(path))
        paragraphs = []
        full_text = []
        sections = []
        current_section = {"title": "Document Start", "content": [], "start_page": 1}

        for para in doc.paragraphs:
            if not para.text.strip():
                continue

            paragraphs.append({
                "text": para.text,
                "style": para.style.name if para.style else None,
                "is_heading": para.style and "heading" in para.style.name.lower() if para.style else False,
            })
            full_text.append(para.text)

            # Detect headings to create sections
            if para.style and "heading" in para.style.name.lower():
                # Save previous section
                if current_section["content"]:
                    current_section["content"] = "\n".join(current_section["content"])
                    sections.append(current_section)
                # Start new section
                current_section = {
                    "title": para.text,
                    "content": [],
                    "start_page": len(sections) + 1,
                    "end_page": len(sections) + 1,
                }
            else:
                current_section["content"].append(para.text)

        # Save last section
        if current_section["content"]:
            current_section["content"] = "\n".join(current_section["content"])
            sections.append(current_section)

        # Extract tables
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)

        combined_text = "\n\n".join(full_text)

        return {
            "filename": path.name,
            "type": "docx",
            "paragraph_count": len(paragraphs),
            "paragraphs": paragraphs,
            "tables": tables,
            "full_text": combined_text,
            "sections": sections,
            "parser": "python-docx",
        }


document_parser = DocumentParser()
